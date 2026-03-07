"""
로컬 파일 로더(FileLoader)

로컬에 저장된 파일(.txt, .pdf, .docx, .json, .csv 등)을 읽어 `Document` 객체로 반환합니다.
간단한 파일 확장자 기반 라우팅을 제공하며, PDF/DOCX는 외부 라이브러리를 사용하여 내용을 추출합니다.
"""

# 이해가게 개념설명 알기쉽게 핵심 설명(자세 설명):
# 목적:
# - 로컬에 있는 여러 형식의 파일을 읽고 내부 텍스트를 추출해 `Document`로 만들어줘요.
# 주요 포인트:
# - 확장자(.pdf, .docx 등)에 따라 적절한 읽기 함수를 호출합니다.
# - 필요하면 외부 라이브러리(PyPDF2, python-docx)를 사용해 텍스트를 추출해요.

from pathlib import Path
from typing import Any

from .base_crawler import BaseCrawler, Document


class FileLoader(BaseCrawler):
    """로컬 파일을 읽어 `Document` 객체로 변환하는 로더."""

    SUPPORTED_EXTENSIONS = {".txt", ".pdf", ".docx", ".md", ".json", ".csv"}

    def __init__(self):
        """초기화(현재 상태 없음)."""
        pass

    def crawl(self, source: str) -> list[Document]:
        """단일 파일을 읽어 `Document` 리스트(보통 1개)를 반환합니다."""
        path = Path(source)

        if not path.exists():
            print(f"File not found: {source}")
            return []

        if path.suffix.lower() not in self.SUPPORTED_EXTENSIONS:
            print(f"Unsupported file type: {path.suffix}")
            return []

        try:
            content = self._load_file(path)
            metadata = {
                "filename": path.name,
                "extension": path.suffix,
                "size": path.stat().st_size,
                "path": str(path.absolute()),
            }

            return [Document(content=content, metadata=metadata, source=source)]

        except Exception as e:
            print(f"Error loading {source}: {e}")
            return []

    def crawl_multiple(self, sources: list[str]) -> list[Document]:
        """여러 파일을 순회하여 문서 리스트를 반환합니다."""
        documents = []
        for source in sources:
            documents.extend(self.crawl(source))
        return documents

    def _load_file(self, path: Path) -> str:
        """확장자에 따라 파일을 읽는 내부 헬퍼."""
        suffix = path.suffix.lower()

        if suffix == ".txt" or suffix == ".md":
            return path.read_text(encoding="utf-8")

        elif suffix == ".pdf":
            return self._load_pdf(path)

        elif suffix == ".docx":
            return self._load_docx(path)

        elif suffix == ".json":
            return path.read_text(encoding="utf-8")

        elif suffix == ".csv":
            return path.read_text(encoding="utf-8")

        return ""

    def _load_pdf(self, path: Path) -> str:
        """PDF 파일의 텍스트를 추출합니다. `PyPDF2` 사용 권장."""
        try:
            import PyPDF2

            with open(path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                text = ""
                for page in reader.pages:
                    text += page.extract_text() + "\n"
                return text
        except ImportError:
            print("PyPDF2 not installed. Run: pip install PyPDF2")
            return ""

    def _load_docx(self, path: Path) -> str:
        """DOCX 파일의 텍스트를 추출합니다. `python-docx` 사용 권장."""
        try:
            from docx import Document as DocxDocument

            doc = DocxDocument(path)
            text = "\n".join([para.text for para in doc.paragraphs])
            return text
        except ImportError:
            print("python-docx not installed. Run: pip install python-docx")
            return ""
